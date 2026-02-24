import os
import json
import requests
import sqlite3
import secrets
import random
import uuid
import time
import re
import base64
import zipfile
import threading
import xml.etree.ElementTree as ET
from html import unescape
from datetime import datetime, timedelta
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_from_directory
from functools import wraps
from werkzeug.utils import secure_filename

# Document text extraction libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[WARNING] PyPDF2 not available. PDF text extraction disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARNING] python-docx not available. DOCX text extraction disabled.")

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False
    print("[WARNING] striprtf not available. RTF text extraction disabled.")

app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'supersecretkey_change_this_in_production')

# Настройка сессий для работы с HTTPS
app.config['SESSION_COOKIE_SECURE'] = True  # True для HTTPS в продакшене
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # Lax позволяет работать с редиректами между доменами
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)

# Configuration
API_URL = "https://api.intelligence.io.solutions/api/v1/chat/completions"
API_TOKEN = "io-v2-eyJhbGciOiJSUzI1NiIsInR5cCI6IkpXVCJ9.eyJvd25lciI6IjY0ODczZDQ3LWQxYzMtNDA5My1iMDAyLTY4YWFiZmQ5YjJjNSIsImV4cCI6NDkxODQ3NjYwNH0.iAwhAprPStFvgrktcGEmvd5J3W7S2o6QxNwt0H2eVTZQxnV-ZE1FUfA5oQ7QJZAyTozsxUYwnIKTsI7PwkUecA"
DOWNLOADER_API_URL = "https://download.dreampartners.online"
DB_FILE = 'chat_history.db'
UPLOAD_FOLDER = 'static/uploads/avatars'
CHAT_IMAGES_FOLDER = 'static/uploads/chat_images'
CHAT_DOCUMENTS_FOLDER = 'static/uploads/chat_documents'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc', 'rtf', 'odt'}
ALLOWED_VIDEO_EXTENSIONS = {'mp4', 'mov', 'avi', 'mkv', 'webm'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CHAT_IMAGES_FOLDER, exist_ok=True)
os.makedirs(CHAT_DOCUMENTS_FOLDER, exist_ok=True)

# Main system prompt (hidden from users, controls censorship, behavior, etc.)
MAIN_SYSTEM_PROMPT = "Отвечай кратко и по делу. Избегай лишних слов и длинных объяснений, если не требуется детальный ответ."

# Special system prompt for dreamGPT AI
DREAMGPT_SYSTEM_PROMPT = "Ты умный ассистент dreamGPT AI. Отвечай кратко, ясно и эффективно."

# SSO Configuration
SSO_AUTH_URL = "https://auth.dreampartners.online"
SSO_CLIENT_ID = "dream_ai"
SSO_CLIENT_SECRET = os.environ.get('SSO_CLIENT_SECRET', '')
# Production redirect URI (registered in dreamID)
SSO_REDIRECT_URI = "https://ai.dreampartners.online/callback" 

# Models with friendly names and descriptions
MODELS_INFO = {
    'dreamgpt-ai': {
        'name': 'dreamGPT AI',
        'description': 'Умная модель: автоматически выбирает GPT для текста и Qwen Vision для изображений',
        'best_for': ['Универсальные задачи', 'Работа с текстом и изображениями', 'Автоматический выбор'],
        'rating': 5,
        'is_smart': True
    },
    'openai/gpt-oss-120b': {
        'name': 'GPT-4 OSS',
        'description': 'Мощная модель для сложных задач и анализа',
        'best_for': ['Анализ данных', 'Сложные рассуждения', 'Научные вопросы'],
        'rating': 5
    },
    'deepseek-ai/DeepSeek-V3.2': {
        'name': 'DeepSeek V3',
        'description': 'Универсальный помощник с отличным пониманием контекста',
        'best_for': ['Разговоры', 'Объяснения', 'Творческие задачи'],
        'rating': 5
    },
    'mistralai/Mistral-Large-Instruct-2411': {
        'name': 'Mistral Large',
        'description': 'Быстрая и точная модель для ежедневных задач',
        'best_for': ['Быстрые ответы', 'Письмо', 'Кодинг'],
        'rating': 5
    },
    'moonshotai/Kimi-K2-Thinking': {
        'name': 'Kimi K2',
        'description': 'Модель с расширенным мышлением для глубокого анализа',
        'best_for': ['Анализ проблем', 'Планирование', 'Логика'],
        'rating': 4
    },
    'Qwen/Qwen3-Next-80B-A3B-Instruct': {
        'name': 'Qwen Next',
        'description': 'Продвинутая модель для профессиональных задач',
        'best_for': ['Профессиональные задачи', 'Многоязычность', 'Технические вопросы'],
        'rating': 4
    },
    'meta-llama/Llama-3.3-70B-Instruct': {
        'name': 'Llama 3.3',
        'description': 'Сбалансированная модель для разных задач',
        'best_for': ['Общие задачи', 'Обучение', 'Исследования'],
        'rating': 4
    },
    'deepseek-ai/DeepSeek-R1-0528': {
        'name': 'DeepSeek R1',
        'description': 'Модель с рефлексией для точных ответов',
        'best_for': ['Точность', 'Детальный анализ', 'Сложные вопросы'],
        'rating': 4
    },
    'mistralai/Devstral-Small-2505': {
        'name': 'Devstral',
        'description': 'Специализированная модель для разработки',
        'best_for': ['Программирование', 'Отладка', 'Техническая документация'],
        'rating': 4
    },
    'Qwen/Qwen2.5-VL-32B-Instruct': {
        'name': 'Qwen Vision',
        'description': 'Модель с поддержкой изображений',
        'best_for': ['Работа с изображениями', 'Визуальный анализ'],
        'rating': 3
    },
    'zai-org/GLM-4.6': {
        'name': 'GLM 4.6',
        'description': 'Универсальный помощник',
        'best_for': ['Общие вопросы', 'Быстрые задачи'],
        'rating': 3
    }
}

# Default models list (dreamGPT AI first)
MODELS = ['dreamgpt-ai'] + [m for m in MODELS_INFO.keys() if m != 'dreamgpt-ai']

# Top 3 recommended models
TOP_MODELS = ['dreamgpt-ai', 'openai/gpt-oss-120b', 'deepseek-ai/DeepSeek-V3.2']

# Models that support vision/images (based on actual capabilities)
VISION_MODELS = [
    'dreamgpt-ai',  # Smart model that auto-selects vision model for images
    'Qwen/Qwen2.5-VL-32B-Instruct',  # Explicitly supports vision
    # Add other vision models here when confirmed
]

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_clean_text_from_html(html_content):
    """Extract readable text from HTML: drop scripts/styles/tags, unescape, normalize spaces."""
    if not html_content:
        return ""

    # Remove script/style blocks
    html_content = re.sub(r'<script[^>]*>.*?</script>', ' ', html_content, flags=re.DOTALL | re.IGNORECASE)
    html_content = re.sub(r'<style[^>]*>.*?</style>', ' ', html_content, flags=re.DOTALL | re.IGNORECASE)

    # Strip all tags
    html_content = re.sub(r'<[^>]+>', ' ', html_content)

    # Unescape HTML entities
    html_content = unescape(html_content)

    # Normalize whitespace
    html_content = re.sub(r'\s+', ' ', html_content).strip()

    return html_content


def fetch_url_text(url):
    """Fetch URL and return clean text; returns (text, error_message)."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                          "(KHTML, like Gecko) Chrome/124.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "ru,en;q=0.9",
        }
        resp = requests.get(url, headers=headers, timeout=12, allow_redirects=True)
        if resp.status_code != 200:
            return None, f"Не удалось загрузить сайт (код {resp.status_code})."

        content_type = resp.headers.get('Content-Type', '')
        # Limit body size to prevent huge downloads
        body = resp.content
        max_bytes = 800_000  # ~800KB
        if len(body) > max_bytes:
            body = body[:max_bytes]

        if 'text/html' in content_type.lower() or '<html' in (resp.text[:200].lower() if resp.text else ''):
            text = extract_clean_text_from_html(body.decode(resp.encoding or 'utf-8', errors='ignore'))
        elif 'text/plain' in content_type.lower():
            text = body.decode(resp.encoding or 'utf-8', errors='ignore')
            text = re.sub(r'\s+', ' ', text).strip()
        else:
            return None, "Сайт вернул неподдерживаемый формат (не HTML/текст)."

        if not text or len(text) < 20:
            return None, "Не удалось извлечь содержимое сайта или оно слишком короткое."

        # Truncate to keep prompt concise
        if len(text) > 5000:
            text = text[:5000] + " ..."

        return text, None
    except requests.exceptions.RequestException as e:
        return None, f"Ошибка при загрузке сайта: {e}"
    except Exception as e:
        return None, f"Неожиданная ошибка при обработке сайта: {e}"

def allowed_document_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_DOCUMENT_EXTENSIONS

def allowed_video_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_VIDEO_EXTENSIONS

def image_url_to_base64(image_url):
    """Convert local image URL to base64 data URI for API"""
    try:
        # Remove leading slash if present
        if image_url.startswith('/'):
            image_url = image_url[1:]
        
        # Check if it's a chat image
        if image_url.startswith('static/uploads/chat_images/'):
            filepath = image_url
        else:
            return None
        
        # Read file and convert to base64
        if os.path.exists(filepath):
            with open(filepath, 'rb') as f:
                image_data = f.read()
                base64_data = base64.b64encode(image_data).decode('utf-8')
                
                # Determine MIME type from extension
                ext = filepath.rsplit('.', 1)[1].lower() if '.' in filepath else 'png'
                mime_types = {
                    'png': 'image/png',
                    'jpg': 'image/jpeg',
                    'jpeg': 'image/jpeg',
                    'gif': 'image/gif',
                    'webp': 'image/webp'
                }
                mime_type = mime_types.get(ext, 'image/png')
                
                return {
                    'type': mime_type,
                    'data': base64_data
                }
    except Exception as e:
        print(f"Error converting image to base64: {e}")
    return None

def init_db():
    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.cursor()
        # Table for chat sessions
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS chats (
                id TEXT PRIMARY KEY,
                user_id TEXT,
                title TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                model TEXT
            )
        ''')
        # Table for messages
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT,
                role TEXT,
                content TEXT,
                images_json TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(chat_id) REFERENCES chats(id)
            )
        ''')
        # Add images_json column if it doesn't exist
        try:
            cursor.execute('ALTER TABLE messages ADD COLUMN images_json TEXT')
        except sqlite3.OperationalError:
            pass  # Column already exists
        # Add documents_json column if it doesn't exist
        try:
            cursor.execute('ALTER TABLE messages ADD COLUMN documents_json TEXT')
        except sqlite3.OperationalError:
            pass  # Column already exists
        # Table for user profiles
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_profiles (
                user_id TEXT PRIMARY KEY,
                custom_name TEXT,
                avatar_url TEXT,
                system_prompt TEXT,
                preferred_model TEXT,
                prompt_template TEXT,
                memory_enabled INTEGER DEFAULT 1,
                notifications_enabled INTEGER DEFAULT 0,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        # Add memory_enabled and notifications_enabled columns if they don't exist
        try:
            cursor.execute('ALTER TABLE user_profiles ADD COLUMN memory_enabled INTEGER DEFAULT 1')
        except sqlite3.OperationalError:
            pass
        try:
            cursor.execute('ALTER TABLE user_profiles ADD COLUMN notifications_enabled INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
        # Table for user memory (important facts to remember)
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS user_memory (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT NOT NULL,
                memory_text TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(user_id) REFERENCES user_profiles(user_id)
            )
        ''')
        # Table for welcome topics (cached AI-generated topics)
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS welcome_topics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT,
                description TEXT,
                prompt TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()

init_db()

# Helper functions
def clean_reasoning_tags(text):
    """Remove all reasoning/thinking tags from AI responses"""
    if not text or not isinstance(text, str):
        return text
    
    # Remove common reasoning tags with regex first (most efficient)
    # Remove all variations of reasoning tags
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<thinking>.*?</thinking>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<reasoning>.*?</reasoning>', '', text, flags=re.DOTALL | re.IGNORECASE)
    
    # Remove lines that start with reasoning markers (line-by-line processing for edge cases)
    lines = text.split('\n')
    cleaned_lines = []
    in_reasoning_block = False
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if we're entering a reasoning block
        if '<think' in line_lower or '<reasoning' in line_lower or '<redacted_reasoning' in line_lower:
            in_reasoning_block = True
            continue
        
        # Check if we're exiting a reasoning block
        if in_reasoning_block and ('</think>' in line_lower or '</reasoning>' in line_lower or '</think>' in line_lower):
            in_reasoning_block = False
            continue
        
        # Skip lines inside reasoning blocks
        if in_reasoning_block:
            continue
        
        # Skip lines that are just reasoning markers
        if line_lower.startswith('<') and ('think' in line_lower or 'reasoning' in line_lower):
            continue
        
        cleaned_lines.append(line)
    
    result = '\n'.join(cleaned_lines).strip()
    
    # Remove any remaining reasoning markers at start/end
    result = re.sub(r'^<[^>]*?(?:think|reasoning)[^>]*?>.*?</[^>]*?(?:think|reasoning)[^>]*?>', '', result, flags=re.DOTALL | re.IGNORECASE)
    result = re.sub(r'^\[REASONING\].*?\[/REASONING\]', '', result, flags=re.DOTALL | re.IGNORECASE)
    
    # Clean up multiple newlines
    result = re.sub(r'\n{3,}', '\n\n', result)
    
    return result.strip()

def get_user_memory(user_id):
    """Get all memory entries for a user as formatted text for system prompt"""
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT memory_text FROM user_memory WHERE user_id = ? ORDER BY created_at DESC LIMIT 20', (user_id,))
            rows = cursor.fetchall()
            if not rows:
                return None
            memory_items = [row[0] for row in rows]
            return "ВАЖНАЯ ИНФОРМАЦИЯ О ПОЛЬЗОВАТЕЛЕ (запомни и используй в ответах):\n" + "\n".join([f"- {item}" for item in memory_items])
    except Exception as e:
        print(f"[MEMORY] Error getting user memory: {e}")
        return None

def analyze_and_store_memory(user_id, user_message, ai_response):
    """Automatically analyze conversation and store important facts to memory.
    Returns list of newly added memory entries (text strings)."""
    new_memories = []
    try:
        # Check if memory is enabled for user
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT memory_enabled FROM user_profiles WHERE user_id = ?', (user_id,))
            row = cursor.fetchone()
            if not row or not row[0]:
                return new_memories  # Memory disabled
        
        # Use a fast model to analyze if there's something worth remembering
        # Get existing memory to avoid duplicates
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT memory_text FROM user_memory WHERE user_id = ? ORDER BY created_at DESC LIMIT 10', (user_id,))
            existing_memories = [row[0].lower() for row in cursor.fetchall()]
        
        existing_context = "\n".join([f"- {m}" for m in existing_memories[:5]]) if existing_memories else "Память пуста."
        
        analysis_prompt = f"""Ты анализируешь диалог, чтобы определить, есть ли НОВАЯ важная информация о пользователе для запоминания.

Сообщение пользователя: {user_message[:400]}

Ответ ассистента: {ai_response[:400]}

Уже запомнено:
{existing_context}

Правила:
1. Записывай ТОЛЬКО если:
   - Пользователь ЯВНО сообщил о себе факт (имя, возраст, место жительства, профессия, предпочтения)
   - Пользователь ЯВНО попросил запомнить что-то (например: "запомни это", "запиши", "сохрани")
   - Есть конкретная важная информация о пользователе
2. НЕ записывай:
   - Общие фразы типа "Понял", "Хорошо", "Давайте начнем", "Спасибо", "Окей", "Понятно"
   - Общие советы или информацию из ответа ассистента
   - Вопросы пользователя
   - Короткие подтверждения без смысла
   - Фразы вежливости без фактов
   - Общие рассуждения или мнения без конкретики
   - Информацию, которая не относится к личности пользователя
3. НЕ записывай дубликаты того, что уже запомнено
4. Записывай ТОЛЬКО если есть реальный конкретный факт о пользователе
5. Если НЕТ новой важной информации - верни ТОЛЬКО слово "SKIP" (без кавычек, без точек, без других слов)

Примеры НЕ записывать:
- "Понял, давайте начнем"
- "Хорошо, спасибо"
- "Окей"
- "Понятно"
- "Я программист вообще" (слишком общее)
- Любые общие фразы без конкретики

Примеры записать:
- "Пользователю 25 лет" (конкретный возраст)
- "Пользователь работает программистом в компании X" (конкретная профессия и место)
- "Пользователь живет в Москве" (конкретное место)
- "Пользователь предпочитает Python" (конкретное предпочтение)

Формат ответа: одно короткое предложение с фактом ИЛИ слово "SKIP"."""

        # Use a fast model for analysis
        try:
            response = requests.post(API_URL, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {API_TOKEN}"
            }, json={
                "model": "mistralai/Mistral-Large-Instruct-2411",  # Fast model
                "messages": [
                    {"role": "user", "content": analysis_prompt}
                ],
                "temperature": 0.3,
                "max_tokens": 100
            }, timeout=10)
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    analysis = result['choices'][0]['message'].get('content', '').strip()
                    analysis = clean_reasoning_tags(analysis)
                    
                    # Check if analysis says to skip
                    if not analysis or analysis.upper().strip() in ['SKIP', 'НЕТ', 'NO', 'НЕТ ИНФОРМАЦИИ', 'НЕТ ВАЖНОЙ ИНФОРМАЦИИ']:
                        print(f"[MEMORY] Skipped - no new information to remember")
                        return new_memories  # Return empty list, not None
                    
                    # Additional filter for generic phrases
                    generic_phrases = [
                        'понял', 'давайте начнем', 'хорошо', 'спасибо', 'окей', 'ок', 'понятно',
                        'принял', 'ясно', 'согласен', 'да', 'конечно', 'хорошо, спасибо',
                        'понял, давайте', 'начать', 'начинаем', 'готов', 'готов начать'
                    ]
                    analysis_lower = analysis.lower().strip()
                    if any(phrase in analysis_lower for phrase in generic_phrases):
                        print(f"[MEMORY] Skipped - generic phrase detected: {analysis}")
                        return new_memories
                    
                    # Check if analysis is too short or doesn't contain factual information
                    if len(analysis.strip()) < 10:
                        print(f"[MEMORY] Skipped - too short: {analysis}")
                        return new_memories
                    
                    # Normalize for comparison
                    analysis_normalized = analysis_lower
                    
                    # More strict duplicate check
                    with sqlite3.connect(DB_FILE) as conn:
                        cursor = conn.cursor()
                        cursor.execute('SELECT memory_text FROM user_memory WHERE user_id = ?', (user_id,))
                        existing = cursor.fetchall()
                        
                        # Check for duplicates - more strict matching
                        is_duplicate = False
                        for (existing_text,) in existing:
                            existing_normalized = existing_text.lower().strip()
                            
                            # Exact match
                            if analysis_normalized == existing_normalized:
                                is_duplicate = True
                                break
                            
                            # Check if one contains the other (for similar facts)
                            # Only if both are short enough (to avoid false positives)
                            if len(analysis_normalized) < 100 and len(existing_normalized) < 100:
                                # Check key words overlap
                                analysis_words = set(analysis_normalized.split())
                                existing_words = set(existing_normalized.split())
                                
                                # If more than 50% words overlap, consider duplicate
                                if len(analysis_words) > 0 and len(existing_words) > 0:
                                    overlap = len(analysis_words & existing_words)
                                    overlap_ratio = overlap / max(len(analysis_words), len(existing_words))
                                    if overlap_ratio > 0.5:
                                        is_duplicate = True
                                        break
                        
                        if not is_duplicate and len(analysis.strip()) > 5:  # Minimum length check
                            cursor.execute('INSERT INTO user_memory (user_id, memory_text) VALUES (?, ?)', (user_id, analysis.strip()))
                            memory_id = cursor.lastrowid
                            conn.commit()
                            print(f"[MEMORY] Stored new memory: {analysis}")
                            new_memories.append({
                                'id': memory_id,
                                'text': analysis.strip()
                            })
                        else:
                            print(f"[MEMORY] Skipped duplicate or too short: {analysis}")
        except Exception as e:
            print(f"[MEMORY] Error analyzing memory: {e}")
    except Exception as e:
        print(f"[MEMORY] Error in analyze_and_store_memory: {e}")
    
    return new_memories

def generate_chat_title(user_message, images_info=None):
    """Generate a short, descriptive chat title using AI - fast model, no system prompts"""
    message_text = user_message if isinstance(user_message, str) else ""
    
    # If no message but has images, use placeholder
    if (not message_text or not message_text.strip()) and images_info and len(images_info) > 0:
        message_text = "Изображение"
    
    if not message_text or not message_text.strip():
        return None
    
    try:
        print(f"[TITLE GEN] Calling API for title generation with message: '{message_text[:50]}'")
        # Simple prompt - if has images, mention it
        if images_info and len(images_info) > 0:
            prompt = f"""Ответь кратким названием (2-4 слова) для чата если первое сообщение с фото: "{message_text[:150]}"

Только название на русском, без кавычек."""
        else:
            prompt = f"""Ответь кратким названием (2-4 слова) для чата если первое сообщение: "{message_text[:150]}"

Только название на русском, без кавычек."""
        
        # Use fast NON-REASONING models only - separate parallel request
        # Avoid reasoning/thinking models to prevent reasoning tags in output
        fast_models = [
            "mistralai/Mistral-Large-Instruct-2411",  # Fast, no reasoning
            "deepseek-ai/DeepSeek-V3.2",  # Good balance, no reasoning
            "meta-llama/Llama-3.3-70B-Instruct",  # Fast, no reasoning
            "Qwen/Qwen3-Next-80B-A3B-Instruct",  # Alternative
            MODELS[0]  # Fallback to first available
        ]
        
        last_error = None
        for model_name in fast_models:
            try:
                print(f"[TITLE GEN] Sending request to API with model: {model_name}")
                response = requests.post(API_URL, headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {API_TOKEN}"
                }, json={
                    "model": model_name,
                    "messages": [
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.5,
                    "max_tokens": 20
                }, timeout=10)
                
                print(f"[TITLE GEN] API response status: {response.status_code}")
                if response.status_code == 200:
                    result = response.json()
                    print(f"[TITLE GEN] API response: {result}")
                    if 'choices' in result and len(result['choices']) > 0:
                        content = result['choices'][0]['message'].get('content')
                        print(f"[TITLE GEN] Extracted content: '{content}'")
                        if content and isinstance(content, str):
                            # Remove reasoning tags completely
                            title = clean_reasoning_tags(content)
                            title = title.strip().strip('"').strip("'").strip('«').strip('»')
                            # Remove any extra text after title - take only first line
                            title = title.split('\n')[0].strip()
                            # Remove any markdown formatting
                            title = re.sub(r'^\*\*', '', title)
                            title = re.sub(r'\*\*$', '', title)
                            title = title.strip()
                            # Final validation - must be short and not contain tags
                            if title and len(title) <= 50 and not title.startswith('<') and not '<' in title:
                                print(f"[TITLE GEN] Final title: '{title}'")
                                return title
                    break  # Success, exit loop
                else:
                    last_error = f"Status {response.status_code}: {response.text[:200]}"
                    print(f"[TITLE GEN] Model {model_name} failed: {last_error}")
                    continue  # Try next model
            except requests.exceptions.Timeout:
                print(f"[TITLE GEN] Model {model_name} timeout, trying next...")
                last_error = "Timeout"
                continue
            except Exception as e:
                print(f"[TITLE GEN] Model {model_name} error: {e}")
                last_error = str(e)
                continue
        
        if last_error:
            print(f"[TITLE GEN] All models failed. Last error: {last_error}")
    except Exception as e:
        print(f"[TITLE GEN] Error generating title: {e}")
        import traceback
        traceback.print_exc()
    
    print(f"[TITLE GEN] Returning None - title generation failed")
    return None  # Use fallback

# Pre-generated welcome topics (multiple sets cached) - Updated for 2025
ALL_TOPICS = [
    {'title': 'AI и технологии', 'desc': 'Тренды ИИ в 2025 году', 'prompt': 'Какие главные тренды искусственного интеллекта ожидаются в 2025 году?'},
    {'title': 'Идеи для подарка', 'desc': 'Подарок коллеге', 'prompt': 'Предложи 5 идей для подарка коллеге-программисту'},
    {'title': 'Программирование', 'desc': 'Изучить новый язык', 'prompt': 'Какой язык программирования стоит изучить в 2025 году для начинающего?'},
    {'title': 'Производительность', 'desc': 'Методы тайм-менеджмента', 'prompt': 'Какие современные методы управления временем наиболее эффективны в 2025?'},
    {'title': 'Путешествия', 'desc': 'Планирование отпуска', 'prompt': 'Составь план путешествия на выходные в необычное место'},
    {'title': 'Здоровый образ жизни', 'desc': 'Утренний распорядок', 'prompt': 'Составь идеальный утренний распорядок дня для продуктивности'},
    {'title': 'Финансы', 'desc': 'Инвестиции в 2025', 'prompt': 'Какие инвестиционные стратегии актуальны в 2025 году?'},
    {'title': 'Карьера', 'desc': 'Развитие навыков', 'prompt': 'Какие навыки будут наиболее востребованы на рынке труда в 2025 году?'},
    {'title': 'Образование', 'desc': 'Онлайн-обучение', 'prompt': 'Посоветуй лучшие онлайн-платформы для обучения новым навыкам в 2025'},
    {'title': 'Творчество', 'desc': 'Идеи для контента', 'prompt': 'Придумай 10 интересных идей для постов в блог про технологии и AI'},
    {'title': 'Кулинария', 'desc': 'Быстрый ужин', 'prompt': 'Что приготовить на ужин за 15 минут из простых продуктов?'},
    {'title': 'Отдых', 'desc': 'Фильмы для просмотра', 'prompt': 'Посоветуй 3 интересных фильма или сериала 2024-2025 года'},
    {'title': 'Психология', 'desc': 'Работа со стрессом', 'prompt': 'Дай 5 практических техник для управления стрессом на работе'},
    {'title': 'Наука', 'desc': 'Новые открытия', 'prompt': 'Какие научные открытия 2024-2025 года кажутся наиболее перспективными?'},
    {'title': 'Бизнес', 'desc': 'Стартап-идеи', 'prompt': 'Оцени идею стартапа, связанного с AI-ассистентами в 2025 году'},
]

def init_welcome_topics():
    """Initialize welcome topics in database - generate fresh topics on startup"""
    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.cursor()
        # Clear old topics to ensure fresh ones
        cursor.execute('DELETE FROM welcome_topics')
        
        # Generate multiple sets (all topics shuffled)
        for topic in ALL_TOPICS:
            cursor.execute('''
                INSERT INTO welcome_topics (title, description, prompt)
                VALUES (?, ?, ?)
            ''', (topic['title'], topic['desc'], topic['prompt']))
        conn.commit()

# Initialize topics on startup
init_welcome_topics()

# Decorator for login requirement
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            return redirect(url_for('login_page'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/login')
def login_page():
    return render_template('login.html')

@app.route('/auth/login')
def auth_login():
    # Step 1: Redirect to SSO
    # Generate random state for CSRF protection
    state = secrets.token_urlsafe(16)
    session['oauth_state'] = state
    
    auth_url = f"{SSO_AUTH_URL}/sso?client_id={SSO_CLIENT_ID}&redirect_uri={SSO_REDIRECT_URI}&state={state}"
    return redirect(auth_url)

@app.route('/callback')
def callback():
    code = request.args.get('code')
    state = request.args.get('state')
    error = request.args.get('error')
    
    # Логирование для отладки
    print(f"Callback received: code={'present' if code else 'missing'}, state={state}, error={error}")
    
    # Если есть ошибка от SSO сервера
    if error:
        error_description = request.args.get('error_description', error)
        print(f"SSO Error in callback: {error_description}")
        return f"SSO Error: {error_description}", 400
    
    if not code:
        print("Callback error: No authorization code provided")
        return "No authorization code provided", 400
    
    # Проверка state - если передан, проверяем (защита от CSRF)
    # Но не блокируем, если state отсутствует (код уже одноразовый)
    saved_state = session.get('oauth_state')
    if state and saved_state:
        if state != saved_state:
            print(f"State mismatch: received={state}, saved={saved_state}")
            # Не блокируем - код все равно одноразовый, но логируем
    elif state and not saved_state:
        print(f"State received but not in session (normal after domain redirect): {state}")
    elif not state and saved_state:
        print(f"State expected but not received (may be normal)")

    try:
        # Step 3: Exchange code for token
        token_resp = requests.post(f"{SSO_AUTH_URL}/api/sso/token", json={
            "code": code,
            "client_id": SSO_CLIENT_ID,
            "client_secret": SSO_CLIENT_SECRET
        }, timeout=10)
        
        if token_resp.status_code != 200:
            error_msg = f"Token exchange failed with status {token_resp.status_code}"
            try:
                error_data = token_resp.json()
                error_msg = error_data.get('error', error_msg)
            except:
                error_msg = token_resp.text or error_msg
            print(f"Token exchange error: {error_msg}")
            return f"Login Error: {error_msg}", 400
        
        token_data = token_resp.json()
        
        if 'error' in token_data:
            return f"Login Error: {token_data.get('error')}", 400
            
        access_token = token_data.get('access_token')
        
        if not access_token:
            return "No access token received", 400

        # Step 4: Get user info
        user_resp = requests.get(f"{SSO_AUTH_URL}/api/sso/user", headers={
            "Authorization": f"Bearer {access_token}"
        }, timeout=10)
        
        if user_resp.status_code != 200:
            error_msg = f"Failed to get user info: {user_resp.status_code}"
            try:
                error_data = user_resp.json()
                error_msg = error_data.get('error', error_msg)
            except:
                pass
            return f"Error getting user info: {error_msg}", 400
        
        user_data = user_resp.json()
        
        if 'error' in user_data:
            return f"Error: {user_data.get('error')}", 400
        
        # Store user in session
        session['user'] = user_data
        session.permanent = True
        
        # Save full user data to DB for analytics
        try:
            user_id = str(user_data.get('id'))
            dream_id = user_data.get('sub') or user_data.get('id') # Depending on provider
            email = user_data.get('email')
            name = user_data.get('username') or user_data.get('name')
            avatar = user_data.get('avatar') or user_data.get('picture')
            language = user_data.get('locale') or user_data.get('language') or 'ru'
            
            # Get UTM source from session
            utm_source = session.get('utm_source', 'direct')
            
            full_json_str = json.dumps(user_data, ensure_ascii=False)
            
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                
                # Check if user exists
                cursor.execute('SELECT user_id FROM user_profiles WHERE user_id = ?', (user_id,))
                exists = cursor.fetchone()
                
                if exists:
                    # Update existing user (keep custom settings like system_prompt)
                    cursor.execute('''
                        UPDATE user_profiles 
                        SET dream_id=?, email=?, full_json=?, last_login=CURRENT_TIMESTAMP, language=?, 
                            custom_name = COALESCE(custom_name, ?), 
                            avatar_url = COALESCE(avatar_url, ?)
                        WHERE user_id=?
                    ''', (dream_id, email, full_json_str, language, name, avatar, user_id))
                else:
                    # Create new user
                    cursor.execute('''
                        INSERT INTO user_profiles 
                        (user_id, dream_id, email, custom_name, avatar_url, language, utm_source, full_json, last_login)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP)
                    ''', (user_id, dream_id, email, name, avatar, language, utm_source, full_json_str))
                
                conn.commit()
                print(f"User {user_id} data synced to DB. Source: {utm_source}")
                
        except Exception as db_e:
            print(f"Error saving user to DB: {db_e}")
            import traceback
            traceback.print_exc()
        
        # Очищаем state из сессии (уже использован)
        session.pop('oauth_state', None)
        
        return redirect(url_for('index'))
        
    except requests.exceptions.RequestException as e:
        print(f"Request error during authentication: {e}")
        return f"Network error during authentication: {str(e)}", 500
    except Exception as e:
        print(f"Authentication error: {e}")
        import traceback
        traceback.print_exc()
        return f"Authentication failed: {str(e)}", 500

@app.route('/logout')
def logout():
    session.pop('user', None)
    return redirect(url_for('login_page'))

@app.route('/')
def index():
    # Capture UTM source if present
    utm_source = request.args.get('utm_source')
    if utm_source:
        session['utm_source'] = utm_source
        
    user = session.get('user')
    if not user:
        return render_template('landing.html')
    return render_template('index.html', user=user)

@app.route('/chat/<chat_id>')
@login_required
def chat_page(chat_id):
    """Route for direct chat links - renders index.html which will load the chat"""
    user = session.get('user')
    if not user:
        return redirect(url_for('index'))
    return render_template('index.html', user=user)

@app.route('/api/models')
@login_required
def get_models():
    # Return full model info including top recommendations
    # Also include user's preferred model
    user_id = str(session['user'].get('id'))
    preferred_model = None
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT preferred_model FROM user_profiles WHERE user_id = ?', (user_id,))
            row = cursor.fetchone()
            if row and row[0]:
                preferred_model = row[0]
    except:
        pass
    
    # Return models in correct order (dreamGPT first)
    ordered_models = {}
    # Add dreamGPT first
    if 'dreamgpt-ai' in MODELS_INFO:
        ordered_models['dreamgpt-ai'] = MODELS_INFO['dreamgpt-ai']
    # Add others in MODELS order
    for model_id in MODELS:
        if model_id != 'dreamgpt-ai' and model_id in MODELS_INFO:
            ordered_models[model_id] = MODELS_INFO[model_id]
    
    return jsonify({
        'models': ordered_models,
        'models_order': MODELS,  # Keep order for frontend
        'top': TOP_MODELS,
        'preferred': preferred_model or 'dreamgpt-ai',
        'vision_models': VISION_MODELS
    })

@app.route('/api/chats', methods=['GET'])
@login_required
def get_chats():
    user_id = str(session['user'].get('id')) # Ensure string for DB consistency
    try:
        with sqlite3.connect(DB_FILE) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            # Filter by user_id
            cursor.execute('SELECT * FROM chats WHERE user_id = ? ORDER BY created_at DESC', (user_id,))
            chats = [dict(row) for row in cursor.fetchall()]
        return jsonify(chats)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chats', methods=['POST'])
@login_required
def create_chat():
    data = request.json
    chat_id = data.get('id')
    title = data.get('title', 'Новый чат')
    model = data.get('model', MODELS[0])
    user_id = str(session['user'].get('id'))
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT 1 FROM chats WHERE id = ?', (chat_id,))
            if not cursor.fetchone():
                cursor.execute('INSERT INTO chats (id, user_id, title, model) VALUES (?, ?, ?, ?)', 
                              (chat_id, user_id, title, model))
                conn.commit()
        return jsonify({'status': 'success', 'id': chat_id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chats/<chat_id>', methods=['DELETE'])
@login_required
def delete_chat(chat_id):
    user_id = str(session['user'].get('id'))
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            # Verify ownership before delete
            cursor.execute('SELECT 1 FROM chats WHERE id = ? AND user_id = ?', (chat_id, user_id))
            if not cursor.fetchone():
                return jsonify({'error': 'Chat not found or access denied'}), 404
                
            cursor.execute('DELETE FROM messages WHERE chat_id = ?', (chat_id,))
            cursor.execute('DELETE FROM chats WHERE id = ?', (chat_id,))
            conn.commit()
        return jsonify({'status': 'success', 'message': 'Chat deleted'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chats/<chat_id>/messages', methods=['GET'])
@login_required
def get_messages(chat_id):
    user_id = str(session['user'].get('id'))
    try:
        with sqlite3.connect(DB_FILE) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            # Verify ownership
            cursor.execute('SELECT 1 FROM chats WHERE id = ? AND user_id = ?', (chat_id, user_id))
            if not cursor.fetchone():
                 # If chat doesn't exist yet (unsaved new chat), allow empty return
                 cursor.execute('SELECT 1 FROM chats WHERE id = ?', (chat_id,))
                 if cursor.fetchone():
                     return jsonify({'error': 'Access denied'}), 403
                 return jsonify([]) # Empty for new chat
                 
            cursor.execute('SELECT id, chat_id, role, content, images_json, documents_json, created_at FROM messages WHERE chat_id = ? ORDER BY id ASC', (chat_id,))
            messages = []
            for row in cursor.fetchall():
                msg = dict(row)
                # Parse images if exists
                if msg.get('images_json'):
                    try:
                        msg['images'] = json.loads(msg['images_json'])
                    except:
                        msg['images'] = []
                else:
                    msg['images'] = []
                # Parse documents if exists
                if msg.get('documents_json'):
                    try:
                        msg['documents'] = json.loads(msg['documents_json'])
                    except:
                        msg['documents'] = []
                else:
                    msg['documents'] = []
                messages.append(msg)
        return jsonify(messages)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def process_video_task(chat_id, video_url_or_path, is_file, user_id, user_message_text, model='dreamgpt-ai'):
    """Background task to process video using external API and then generate AI response"""
    try:
        print(f"[VIDEO] Starting background processing for chat {chat_id}")
        
        # Handle dreamGPT AI - auto-select model based on content
        actual_model = model
        if model == 'dreamgpt-ai':
            # For video processing, use GPT for text (no images in video processing)
            actual_model = 'openai/gpt-oss-120b'
        
        # Step 1: Call Downloader/Transcribe API
        transcription = ""
        summary = ""
        api_endpoint = f"{DOWNLOADER_API_URL}/api/process"
        
        payload = {}
        files = None
        
        if is_file:
            # If it's a local file path, we need to upload it
            # Use /transcribe endpoint which accepts files
            api_endpoint = f"{DOWNLOADER_API_URL}/transcribe"
            if os.path.exists(video_url_or_path):
                files = {'file': open(video_url_or_path, 'rb')}
            else:
                print(f"[VIDEO] File not found: {video_url_or_path}")
                # Save error message to DB
                with sqlite3.connect(DB_FILE) as conn:
                    cursor = conn.cursor()
                    cursor.execute('UPDATE messages SET content = ? WHERE chat_id = ? AND content LIKE ?', 
                                   (f"Ошибка: Файл видео не найден: {video_url_or_path}", chat_id, "%[VIDEO_PROCESSING]%"))
                    conn.commit()
                return
        else:
            # URL
            payload = {"url": video_url_or_path}
        
        try:
            print(f"[VIDEO] Calling external API: {api_endpoint}")
            print(f"[VIDEO] Payload: {payload if not files else 'FILE UPLOAD'}")
            print(f"[VIDEO] URL: {video_url_or_path}")
            
            if files:
                resp = requests.post(api_endpoint, files=files, timeout=300)
                if files['file']: files['file'].close()
            else:
                resp = requests.post(api_endpoint, json=payload, timeout=300)
            
            print(f"[VIDEO] API Response Status: {resp.status_code}")
            print(f"[VIDEO] API Response Headers: {dict(resp.headers)}")
            print(f"[VIDEO] API Response Text (first 500 chars): {resp.text[:500]}")
            
            # Try to parse JSON
            try:
                data = resp.json()
                print(f"[VIDEO] API Response JSON parsed successfully")
            except Exception as json_err:
                print(f"[VIDEO] Failed to parse JSON response: {json_err}")
                print(f"[VIDEO] Full response text: {resp.text}")
                transcription = f"Ошибка: API вернул некорректный ответ (не JSON). Статус: {resp.status_code}"
                data = None
            
            if data and data.get('status') == 'success':
                transcription = data.get('transcription') or data.get('text', '')
                summary = data.get('summary', '')
                print(f"[VIDEO] API success. Transcription len: {len(transcription)}")
            elif data:
                print(f"[VIDEO] API error response: {data}")
                error_msg = data.get('message', 'Не удалось расшифровать видео')
                transcription = f"Ошибка: {error_msg}"
            # else: transcription already set above if JSON parsing failed
        except Exception as e:
            print(f"[VIDEO] API request failed: {e}")
            import traceback
            traceback.print_exc()
            transcription = f"Ошибка при обработке видео: {str(e)}"

        # Step 2: Generate AI response with the video context
        system_prompt = MAIN_SYSTEM_PROMPT.strip()
        
        # Add dreamGPT specific prompt if using dreamGPT AI
        if model == 'dreamgpt-ai':
            if system_prompt:
                system_prompt += "\n\n" + DREAMGPT_SYSTEM_PROMPT
            else:
                system_prompt = DREAMGPT_SYSTEM_PROMPT
        
        # Check if transcription contains an error
        if transcription.startswith("Ошибка:"):
            # Save error message to DB and return
            error_content = transcription
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT id FROM messages WHERE chat_id = ? AND role = 'assistant' AND content LIKE '%[VIDEO_PROCESSING]%' ORDER BY id DESC LIMIT 1", (chat_id,))
                row = cursor.fetchone()
                if row:
                    cursor.execute('UPDATE messages SET content = ? WHERE id = ?', (error_content, row[0]))
                else:
                    cursor.execute('INSERT INTO messages (chat_id, role, content) VALUES (?, ?, ?)', (chat_id, 'assistant', error_content))
                conn.commit()
            return
        
        # Add video context with safety instructions (audio transcript may be imperfect)
        video_context = (
            "\n\n[ВИДЕО КОНТЕКСТ]\n"
            "Пользователь отправил видео по ссылке. Ниже аудио-транскрипт видео (может содержать ошибки распознавания речи).\n"
            f"{transcription}\n"
        )
        if summary:
            video_context += f"Краткое содержание:\n{summary}\n"
        video_context += (
            "[КОНЕЦ ВИДЕО КОНТЕКСТ]\n\n"
            "ВАЖНЫЕ ИНСТРУКЦИИ ДЛЯ ОТВЕТА:\n"
            "1. Начни ответ с фразы: 'Вы отправили видео. Вот что в нем:'\n"
            "2. Затем опиши содержание видео на основе транскрипта\n"
            "3. Помни: транскрипт может содержать ошибки распознавания речи, но НЕ УПОМИНАЙ это в ответе пользователю\n"
            "4. Будь кратким и по делу, избегай домыслов\n"
            "5. Просто опиши содержание видео естественным языком, без упоминаний о качестве транскрипта\n"
            "6. Если транскрипт совсем непонятный или пустой, скажи что не удалось распознать содержание\n"
            "7. ОБЯЗАТЕЛЬНО начни с 'Вы отправили видео. Вот что в нем:' чтобы было понятно, что это расшифровка видео\n"
            "8. НЕ пиши про ошибки распознавания, артефакты или качество транскрипта - просто опиши содержание\n"
        )
        
        combined_prompt = system_prompt + video_context
        
        # Get user memory
        memory_text = get_user_memory(user_id)
        if memory_text:
            combined_prompt += "\n\n" + memory_text

        # Prepare messages
        messages = [{"role": "system", "content": combined_prompt}]
        
        # Get chat history
        with sqlite3.connect(DB_FILE) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute('SELECT role, content FROM messages WHERE chat_id = ? ORDER BY id ASC', (chat_id,))
            rows = cursor.fetchall()
            for row in rows:
                # Skip the placeholder message we added earlier (detected by content)
                content = row['content'] or ''
                if "[VIDEO_PROCESSING]" in content:
                    continue
                # Ensure content is a string and not None
                if content:
                    messages.append({"role": row['role'], "content": str(content)})
        
        # Validate messages before sending
        for msg in messages:
            if not isinstance(msg.get('content'), str):
                msg['content'] = str(msg.get('content', ''))
        
        # Call AI
        try:
            print(f"[VIDEO] Calling AI with video context")
            print(f"[VIDEO] Using model: {actual_model}")
            print(f"[VIDEO] Messages count: {len(messages)}")
            
            ai_resp = requests.post(API_URL, headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {API_TOKEN}"
            }, json={
                "model": actual_model, # Use actual model (handles dreamGPT AI)
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 2000
            }, timeout=60)
            
            print(f"[VIDEO] AI API response status: {ai_resp.status_code}")
            
            ai_content = "Извините, не удалось получить ответ от AI."
            if ai_resp.status_code == 200:
                try:
                    ai_result = ai_resp.json()
                    print(f"[VIDEO] AI API response keys: {ai_result.keys() if isinstance(ai_result, dict) else 'Not a dict'}")
                    
                    if 'choices' in ai_result and len(ai_result['choices']) > 0:
                        ai_content = ai_result['choices'][0]['message']['content']
                        ai_content = clean_reasoning_tags(ai_content)
                        print(f"[VIDEO] AI response length: {len(ai_content)}")
                    else:
                        print(f"[VIDEO] No choices in response or empty choices. Response: {ai_result}")
                        ai_content = f"Извините, AI не вернул ответ. Ответ API: {str(ai_result)[:200]}"
                except Exception as json_e:
                    print(f"[VIDEO] Failed to parse AI response JSON: {json_e}")
                    print(f"[VIDEO] Response text: {ai_resp.text[:500]}")
                    ai_content = f"Ошибка парсинга ответа AI: {str(json_e)}"
            else:
                print(f"[VIDEO] AI API returned non-200 status: {ai_resp.status_code}")
                print(f"[VIDEO] Response text: {ai_resp.text[:500]}")
                try:
                    error_data = ai_resp.json()
                    error_msg = error_data.get('error', {}).get('message', ai_resp.text[:200]) if isinstance(error_data, dict) else ai_resp.text[:200]
                    ai_content = f"Ошибка API (код {ai_resp.status_code}): {error_msg}"
                except:
                    ai_content = f"Ошибка API (код {ai_resp.status_code}): {ai_resp.text[:200]}"
            
            # Step 3: Update DB
            # Replace the placeholder message or add new one
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                # Find the placeholder message (most recent assistant message with special tag)
                cursor.execute("SELECT id FROM messages WHERE chat_id = ? AND role = 'assistant' AND content LIKE '%[VIDEO_PROCESSING]%' ORDER BY id DESC LIMIT 1", (chat_id,))
                row = cursor.fetchone()
                
            if row:
                # Update existing placeholder
                cursor.execute('UPDATE messages SET content = ? WHERE id = ?', (ai_content, row[0]))
            else:
                # Insert new message
                cursor.execute('INSERT INTO messages (chat_id, role, content) VALUES (?, ?, ?)', (chat_id, 'assistant', ai_content))
            conn.commit()
            print(f"[VIDEO] Database updated with AI response")

            # Generate chat title (always for video chats, even if not default)
            try:
                cursor.execute('SELECT title FROM chats WHERE id = ?', (chat_id,))
                chat_row = cursor.fetchone()
                current_title = chat_row[0] if chat_row else None
                print(f"[VIDEO TITLE] Current title: '{current_title}'")
                
                # Update title if it's default or temporary video processing title
                default_titles = {'Новый чат', None, '', '🎥 Обработка видео...'}
                if current_title in default_titles:
                    # Use AI response content for title (best quality), fallback to summary/transcript
                    # Extract meaningful part from AI response (skip the "Вы отправили видео" part)
                    title_source = None
                    if ai_content and len(ai_content) > 50:
                        # Try to extract content after "Вот что в нем:"
                        if "Вот что в нем:" in ai_content:
                            content_after = ai_content.split("Вот что в нем:", 1)[1].strip()
                            title_source = content_after[:300]  # First 300 chars of actual content
                        else:
                            title_source = ai_content[:300]
                    
                    # Fallback to summary or transcript if AI response is too short
                    if not title_source or len(title_source) < 20:
                        title_source = summary if summary else (transcription[:300] if transcription else user_message_text)
                    
                    print(f"[VIDEO TITLE] Title source (first 100 chars): '{title_source[:100] if title_source else 'None'}'")
                    
                    if title_source and not title_source.startswith('Ошибка'):
                        ai_title = generate_chat_title(title_source, None)
                        print(f"[VIDEO TITLE] AI generated title: '{ai_title}'")
                        
                        if ai_title and isinstance(ai_title, str) and ai_title.strip() and len(ai_title.strip()) > 2:
                            cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (ai_title, chat_id))
                            conn.commit()
                            print(f"[VIDEO TITLE] ✓ Title set to: '{ai_title}'")
                        else:
                            # Fallback: extract first few words from content
                            print(f"[VIDEO TITLE] AI title generation failed, using fallback")
                            if title_source:
                                words = title_source.split()[:4]
                                fallback_title = ' '.join(words) if words else "🎥 Видео"
                                cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (fallback_title, chat_id))
                                conn.commit()
                                print(f"[VIDEO TITLE] ✓ Fallback title set to: '{fallback_title}'")
                            else:
                                cursor.execute('UPDATE chats SET title = ? WHERE id = ?', ("🎥 Видео", chat_id))
                                conn.commit()
                                print(f"[VIDEO TITLE] ✓ Generic fallback: '🎥 Видео'")
                    else:
                        # Error case - use error title
                        error_title = "❌ Ошибка видео"
                        cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (error_title, chat_id))
                        conn.commit()
                        print(f"[VIDEO TITLE] ✓ Error title set: '{error_title}'")
            except Exception as title_e:
                print(f"[VIDEO TITLE] Exception during title generation: {title_e}")
                import traceback
                traceback.print_exc()
                
        except Exception as e:
            print(f"[VIDEO] AI generation error: {e}")
            import traceback
            traceback.print_exc()
            # Save error message to DB
            error_content = f"Извините, произошла ошибка при генерации ответа: {str(e)}"
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                cursor.execute("SELECT id FROM messages WHERE chat_id = ? AND role = 'assistant' AND content LIKE '%[VIDEO_PROCESSING]%' ORDER BY id DESC LIMIT 1", (chat_id,))
                row = cursor.fetchone()
                if row:
                    cursor.execute('UPDATE messages SET content = ? WHERE id = ?', (error_content, row[0]))
                else:
                    cursor.execute('INSERT INTO messages (chat_id, role, content) VALUES (?, ?, ?)', (chat_id, 'assistant', error_content))
                conn.commit()
            
    except Exception as e:
        print(f"[VIDEO] Task error: {e}")
        import traceback
        traceback.print_exc()
        # Save error message to DB
        error_content = f"Извините, произошла ошибка при обработке видео: {str(e)}"
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id FROM messages WHERE chat_id = ? AND role = 'assistant' AND content LIKE '%[VIDEO_PROCESSING]%' ORDER BY id DESC LIMIT 1", (chat_id,))
            row = cursor.fetchone()
            if row:
                cursor.execute('UPDATE messages SET content = ? WHERE id = ?', (error_content, row[0]))
            else:
                cursor.execute('INSERT INTO messages (chat_id, role, content) VALUES (?, ?, ?)', (chat_id, 'assistant', error_content))
            conn.commit()

@app.route('/api/chat/completions', methods=['POST'])
@login_required
def chat_completions():
    # Ensure threading is available (it's imported at module level)
    import threading
    
    data = request.json
    chat_id = data.get('chat_id')
    user_message = data.get('message', '')  # Allow empty message if images present
    model = data.get('model', MODELS[0])
    images = data.get('images', [])  # Array of image URLs or objects with url
    documents = data.get('documents', [])  # Array of document URLs
    user_id = str(session['user'].get('id'))
    web_context_text = None
    web_context_url = None
    
    if not chat_id:
        return jsonify({'error': 'Missing chat_id'}), 400
    
    # Allow sending only images/documents without text, or text without files, but not empty
    if not user_message and (not images or len(images) == 0) and (not documents or len(documents) == 0):
        return jsonify({'error': 'Missing message, images, or documents'}), 400
    
    # --- VIDEO DETECTION ---
    video_url = None
    video_file_path = None
    is_video_task = False

    # 1. Check for TikTok/YouTube URLs in message
    url_pattern = r'https?://(www\.)?(tiktok\.com|vm\.tiktok\.com|youtube\.com|youtu\.be|instagram\.com)/[^\s]+'
    match = re.search(url_pattern, user_message)
    if match:
        video_url = match.group(0)
        is_video_task = True
        print(f"[VIDEO] Detected video URL: {video_url}")
    else:
        # Detect generic URL (non-video) for website parsing
        generic_url_match = re.search(r'https?://[^\s]+', user_message)
        if generic_url_match:
            candidate_url = generic_url_match.group(0)
            web_context_url = candidate_url
            web_text, web_err = fetch_url_text(candidate_url)
            if web_err:
                # Reply immediately with an error message about the website
                error_msg = f"Извините, не получилось посмотреть сайт по ссылке.\nПричина: {web_err}"
                return jsonify({'role': 'assistant', 'content': error_msg})
            else:
                web_context_text = web_text

    # 2. Check for uploaded video documents
    if documents and not is_video_task:
        for doc in documents:
            doc_url = doc if isinstance(doc, str) else doc.get('url', '')
            if doc_url:
                ext = doc_url.split('.')[-1].lower()
                if ext in ALLOWED_VIDEO_EXTENSIONS:
                     if doc_url.startswith('/static/'):
                         # Convert URL to local path
                         video_file_path = doc_url.lstrip('/')
                         is_video_task = True
                         print(f"[VIDEO] Detected video file: {video_file_path}")
                         break
    
    # Read document contents if any - extract actual text from all supported formats
    document_texts = []
    if documents and len(documents) > 0:
        for doc_url in documents:
            try:
                # Remove leading slash if present
                if isinstance(doc_url, dict): doc_url = doc_url.get('url', '')
                if not doc_url: continue
                
                if doc_url.startswith('/'):
                    doc_url = doc_url[1:]
                
                # Check if it's a chat document
                if doc_url.startswith('static/uploads/chat_documents/'):
                    filepath = doc_url
                    if os.path.exists(filepath):
                        ext = filepath.rsplit('.', 1)[1].lower() if '.' in filepath else ''
                        
                        # Skip video files (handled separately)
                        if ext in ALLOWED_VIDEO_EXTENSIONS:
                            continue
                            
                        filename = os.path.basename(filepath)
                        content = None
                        
                        # Extract text based on file type
                        if ext == 'txt':
                            # Plain text file
                            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                                content = f.read()
                                
                        elif ext == 'pdf' and PDF_AVAILABLE:
                            # PDF file - extract text
                            try:
                                with open(filepath, 'rb') as f:
                                    pdf_reader = PyPDF2.PdfReader(f)
                                    text_parts = []
                                    for page_num, page in enumerate(pdf_reader.pages):
                                        try:
                                            page_text = page.extract_text()
                                            if page_text.strip():
                                                text_parts.append(page_text)
                                        except Exception as page_e:
                                            print(f"[DOCUMENT] Error reading PDF page {page_num}: {page_e}")
                                    if text_parts:
                                        content = '\n\n'.join(text_parts)
                            except Exception as pdf_e:
                                print(f"[DOCUMENT] Error extracting PDF text from {filename}: {pdf_e}")
                                content = None
                                
                        elif ext == 'docx' and DOCX_AVAILABLE:
                            # DOCX file - extract text
                            try:
                                doc = DocxDocument(filepath)
                                paragraphs = []
                                for para in doc.paragraphs:
                                    if para.text.strip():
                                        paragraphs.append(para.text)
                                # Also extract text from tables
                                for table in doc.tables:
                                    for row in table.rows:
                                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                                        if row_text:
                                            paragraphs.append(row_text)
                                if paragraphs:
                                    content = '\n\n'.join(paragraphs)
                            except Exception as docx_e:
                                print(f"[DOCUMENT] Error extracting DOCX text from {filename}: {docx_e}")
                                content = None
                                
                        elif ext == 'rtf' and RTF_AVAILABLE:
                            # RTF file - extract text
                            try:
                                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                                    rtf_content = f.read()
                                    content = rtf_to_text(rtf_content)
                            except Exception as rtf_e:
                                print(f"[DOCUMENT] Error extracting RTF text from {filename}: {rtf_e}")
                                content = None
                                
                        elif ext == 'doc':
                            # Old DOC format - try to read as text (limited support)
                            try:
                                # Try reading as binary and look for text
                                with open(filepath, 'rb') as f:
                                    data = f.read()
                                    # Simple extraction: find readable text chunks
                                    text_chunks = []
                                    current_chunk = []
                                    for byte in data:
                                        if 32 <= byte <= 126 or byte in [9, 10, 13]:  # Printable ASCII
                                            current_chunk.append(chr(byte))
                                        else:
                                            if len(current_chunk) > 10:
                                                text_chunks.append(''.join(current_chunk))
                                            current_chunk = []
                                    if text_chunks:
                                        content = ' '.join(text_chunks)
                            except Exception as doc_e:
                                print(f"[DOCUMENT] Error extracting DOC text from {filename}: {doc_e}")
                                content = None
                                
                        elif ext == 'odt':
                            # ODT file - try to extract as ZIP (ODT is a ZIP archive)
                            try:
                                with zipfile.ZipFile(filepath, 'r') as odt_file:
                                    if 'content.xml' in odt_file.namelist():
                                        xml_content = odt_file.read('content.xml').decode('utf-8', errors='ignore')
                                        # Simple XML text extraction (remove tags)
                                        try:
                                            root = ET.fromstring(xml_content)
                                            # Extract text from all text nodes
                                            text_parts = []
                                            for elem in root.iter():
                                                if elem.text and elem.text.strip():
                                                    text_parts.append(elem.text.strip())
                                            if text_parts:
                                                content = '\n\n'.join(text_parts)
                                        except:
                                            # Fallback: simple regex to remove tags
                                            content = re.sub(r'<[^>]+>', '\n', xml_content)
                                            content = re.sub(r'\n+', '\n', content).strip()
                            except Exception as odt_e:
                                print(f"[DOCUMENT] Error extracting ODT text from {filename}: {odt_e}")
                                content = None
                        
                        # Add extracted content to message
                        if content and content.strip():
                            # Limit content length to avoid token limits (keep first 50000 chars)
                            if len(content) > 50000:
                                content = content[:50000] + "\n\n[... документ обрезан из-за размера ...]"
                            document_texts.append(f"\n\n--- Содержимое документа '{filename}' ---\n{content}\n--- Конец документа '{filename}' ---")
                        else:
                            # If extraction failed, at least mention the file
                            print(f"[DOCUMENT] Could not extract text from {filename} (format: {ext})")
                            document_texts.append(f"\n\nПользователь прикрепил файл '{filename}', но не удалось извлечь текст из него.")
                            
            except Exception as e:
                print(f"[DOCUMENT] Error processing document {doc_url}: {e}")
                import traceback
                traceback.print_exc()
    
    # Normalize empty message to empty string
    if not user_message:
        user_message = ''

    # Get user profile for custom system prompt and name
    with sqlite3.connect(DB_FILE) as conn:
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute('SELECT * FROM user_profiles WHERE user_id = ?', (user_id,))
        profile = cursor.fetchone()
        
        user_name = session['user'].get('username', 'пользователь')
        system_prompt = "Ты полезный ассистент."
        memory_enabled = True  # Default enabled
        if profile:
            if profile['custom_name']:
                user_name = profile['custom_name']
            if profile['system_prompt']:
                system_prompt = profile['system_prompt']
                # Remove {user_name} placeholder - we'll pass name separately
                system_prompt = system_prompt.replace('{user_name}', '')
            if profile['memory_enabled'] is not None:
                memory_enabled = bool(profile['memory_enabled'])
    
    # Handle dreamGPT AI - auto-select model based on content
    actual_model = model
    if model == 'dreamgpt-ai':
        if images and len(images) > 0:
            # Use vision model for images
            actual_model = 'Qwen/Qwen2.5-VL-32B-Instruct'
        else:
            # Use GPT for text
            actual_model = 'openai/gpt-oss-120b'

    with sqlite3.connect(DB_FILE) as conn:
        cursor = conn.cursor()
        
        # Ensure chat exists and belongs to user
        cursor.execute('SELECT user_id FROM chats WHERE id = ?', (chat_id,))
        row = cursor.fetchone()
        
        is_new_chat = False
        if row:
            if row[0] != user_id:
                return jsonify({'error': 'Access denied'}), 403
        else:
            # Create new chat
            is_new_chat = True
            cursor.execute('INSERT INTO chats (id, user_id, title, model) VALUES (?, ?, ?, ?)', 
                          (chat_id, user_id, 'Новый чат', model))

        # Store message with images URLs in DB
        images_json = None
        if images and len(images) > 0:
            # Extract URLs from image objects
            image_urls = []
            for img in images:
                if isinstance(img, str):
                    image_urls.append(img)
                elif isinstance(img, dict) and 'url' in img:
                    image_urls.append(img['url'])
            images_json = json.dumps(image_urls)
        
        # Use placeholder if message is empty but images exist
        message_to_store = user_message if user_message else ('' if not images or len(images) == 0 else '📷 Изображение')
        
        # Store documents JSON
        documents_json = None
        if documents and len(documents) > 0:
            documents_json = json.dumps(documents)
        
        cursor.execute('INSERT INTO messages (chat_id, role, content, images_json, documents_json) VALUES (?, ?, ?, ?, ?)', 
                       (chat_id, 'user', message_to_store, images_json, documents_json))
        conn.commit()

    # --- VIDEO TASK START ---
    if is_video_task:
        # Return "Processing" placeholder immediately
        placeholder_text = "🎥 [VIDEO_PROCESSING] Я смотрю видео, это займет некоторое время... Можете включить уведомления."
        
        # Save placeholder to DB
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('INSERT INTO messages (chat_id, role, content) VALUES (?, ?, ?)', 
                           (chat_id, 'assistant', placeholder_text))
            
            # Set temporary title for new video chats
            if is_new_chat:
                temp_title = "🎥 Обработка видео..."
                cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (temp_title, chat_id))
                print(f"[VIDEO] Set temporary title for new chat: {temp_title}")
            
            conn.commit()
            
        # Start background thread
        target_video = video_url if video_url else video_file_path
        is_file = video_file_path is not None
        
        # Get model from chat or use default
        chat_model = model
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT model FROM chats WHERE id = ?', (chat_id,))
            row = cursor.fetchone()
            if row and row[0]:
                chat_model = row[0]
        
        thread = threading.Thread(target=process_video_task, 
                                args=(chat_id, target_video, is_file, user_id, user_message, chat_model))
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'role': 'assistant', 
            'content': "",
            'processing_video': True
        })
    # -------------------------

    # Prepare external API payload
    # Combine main system prompt with user's custom prompt
    combined_system_prompt = MAIN_SYSTEM_PROMPT.strip()
    
    # Add dreamGPT specific prompt if using dreamGPT AI
    if model == 'dreamgpt-ai':
        if combined_system_prompt:
            combined_system_prompt += "\n\n" + DREAMGPT_SYSTEM_PROMPT
        else:
            combined_system_prompt = DREAMGPT_SYSTEM_PROMPT
    
    if system_prompt and system_prompt.strip():
        combined_system_prompt += "\n\n" + system_prompt.strip()
    
    # Add user memory if enabled
    if memory_enabled:
        memory_text = get_user_memory(user_id)
        if memory_text:
            combined_system_prompt += "\n\n" + memory_text
    
    api_messages = [{"role": "system", "content": combined_system_prompt}]
    
    # Build current user message with images if any
    # Include document texts in API request (but not in displayed message)
    current_user_content = user_message
    if document_texts:
        # Add document content to API request (extract just the content without headers)
        doc_content_for_api = []
        for doc_text in document_texts:
            # Extract content between markers or use the text directly
            if '--- Содержимое документа' in doc_text:
                # Extract content between markers
                parts = doc_text.split('--- Содержимое документа')
                if len(parts) > 1:
                    content_part = parts[1].split('--- Конец документа')[0] if '--- Конец документа' in parts[1] else parts[1]
                    # Remove filename from first line
                    lines = content_part.split('\n', 1)
                    if len(lines) > 1:
                        doc_content_for_api.append(lines[1].strip())
                    else:
                        doc_content_for_api.append(content_part.strip())
                else:
                    doc_content_for_api.append(doc_text.strip())
            else:
                doc_content_for_api.append(doc_text.strip())
        
        if doc_content_for_api:
            doc_text_combined = '\n\n'.join(doc_content_for_api)
            if current_user_content:
                current_user_content = current_user_content + '\n\n' + doc_text_combined
            else:
                current_user_content = doc_text_combined
    
    # Check if model supports vision/images (use actual_model for dreamGPT AI)
    model_supports_vision = actual_model in VISION_MODELS
    
    if images and len(images) > 0:
        if not model_supports_vision:
            return jsonify({'error': f'Модель {model} не поддерживает работу с изображениями. Пожалуйста, выберите модель с поддержкой изображений (например, Qwen Vision).'}), 400
        
        # For vision models, content should be array with text and images
        # If message is empty, use placeholder
        text_part = current_user_content if current_user_content else "Что на фото?"
        content_parts = [{"type": "text", "text": text_part}]
        for img in images:
            # Extract URL from image object or use string directly
            image_url = img if isinstance(img, str) else (img.get('url') or img.get('data'))
            
            if image_url:
                # Convert local image URL to base64 for API
                base64_data = image_url_to_base64(image_url)
                if base64_data:
                    content_parts.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{base64_data['type']};base64,{base64_data['data']}"
                        }
                    })
                else:
                    # Fallback: assume it's already a data URI
                    content_parts.append({
                        "type": "image_url",
                        "image_url": {
                            "url": image_url
                        }
                    })
        current_user_content = content_parts
    
    # Add user name context in first user message (invisible to user)
    with sqlite3.connect(DB_FILE) as conn:
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()
        cursor.execute('SELECT role, content FROM messages WHERE chat_id = ? ORDER BY id ASC', (chat_id,))
        rows = cursor.fetchall()
        is_first_user = True
        for row in rows:
            content = row['content']
            # Add user name context to first user message (invisible to user)
            if row['role'] == 'user' and is_first_user:
                if isinstance(content, str):
                    content = f"[Контекст: пользователя зовут {user_name}] {content}"
                is_first_user = False
            api_messages.append({"role": row['role'], "content": content})
    
    # Add current message with name context if it's first
    if is_new_chat and isinstance(current_user_content, str):
        current_user_content = f"[Контекст: пользователя зовут {user_name}] {current_user_content}"
    elif is_new_chat and isinstance(current_user_content, list) and len(current_user_content) > 0:
        # Add name context to text part
        if current_user_content[0].get('type') == 'text':
            current_user_content[0]['text'] = f"[Контекст: пользователя зовут {user_name}] {current_user_content[0]['text']}"

    # Append web context if available
    if web_context_text and isinstance(current_user_content, str):
        current_user_content = current_user_content + "\n\n[ВЕБ-КОНТЕКСТ]\nПользователь прислал ссылку: " + web_context_url + "\nТекст страницы:\n" + web_context_text + "\n\nВажно: текст мог быть неполным. Не выдавай как дословную цитату, уточни у пользователя при необходимости."
    elif web_context_text and isinstance(current_user_content, list):
        # Prepend to text part if exists, else append new text part
        inserted = False
        for part in current_user_content:
            if part.get('type') == 'text':
                part['text'] = part['text'] + "\n\n[ВЕБ-КОНТЕКСТ]\nПользователь прислал ссылку: " + web_context_url + "\nТекст страницы:\n" + web_context_text + "\n\nВажно: текст мог быть неполным. Не выдавай как дословную цитату, уточни у пользователя при необходимости."
                inserted = True
                break
        if not inserted:
            current_user_content.insert(0, {
                "type": "text",
                "text": f"[ВЕБ-КОНТЕКСТ]\nПользователь прислал ссылку: {web_context_url}\nТекст страницы:\n{web_context_text}\n\nВажно: текст мог быть неполным. Не выдавай как дословную цитату, уточни у пользователя при необходимости."
            })
    
    api_messages.append({"role": "user", "content": current_user_content})

    payload = {
        "model": actual_model,  # Use actual_model (handles dreamGPT AI)
        "messages": api_messages,
        "temperature": 0.7,
        "max_tokens": 2000
    }

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_TOKEN}"
    }

    # For new chats, set temporary title first
    if is_new_chat:
        print(f"[TITLE GEN] New chat detected: {chat_id}, message: '{user_message[:50] if user_message else 'empty'}', images: {len(images) if images else 0}")
        # Set temporary title first
        if user_message and user_message.strip():
            words = user_message.strip().split()
            fallback_title = words[0] if words else "Новый чат"
            if len(words) > 1:
                fallback_title = " ".join(words[:2])
        elif images and len(images) > 0:
            fallback_title = "Изображение"
        else:
            fallback_title = "Новый чат"
        
        print(f"[TITLE GEN] Setting temporary fallback title: '{fallback_title}'")
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (fallback_title, chat_id))
            conn.commit()
        
        # Generate title in parallel IMMEDIATELY - but ONLY if NO images
        # For chats with images, we'll generate title AFTER AI response
        if not images or len(images) == 0:
            import threading
            def generate_title_parallel():
                try:
                    print(f"[TITLE GEN] Starting title generation for chat {chat_id}")
                    title_message = user_message if user_message else ""
                    print(f"[TITLE GEN] Message for title: '{title_message}'")
                    
                    ai_title = generate_chat_title(title_message, None)
                    print(f"[TITLE GEN] Generated title: '{ai_title}'")
                    
                    if not ai_title or not isinstance(ai_title, str) or len(ai_title.strip()) == 0:
                        print(f"[TITLE GEN] Title is empty, using fallback")
                        return  # Use fallback
                    
                    with sqlite3.connect(DB_FILE) as conn:
                        cursor = conn.cursor()
                        cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (ai_title, chat_id))
                        conn.commit()
                        print(f"[TITLE GEN] Title updated to: '{ai_title}' for chat {chat_id}")
                except Exception as e:
                    print(f"[TITLE GEN] Parallel title generation error: {e}")
                    import traceback
                    traceback.print_exc()
            
            # Start title generation IMMEDIATELY in background (only for text-only chats)
            print(f"[TITLE GEN] Launching parallel thread for chat {chat_id} (text-only)")
            threading.Thread(target=generate_title_parallel, daemon=True).start()

    try:
        response = requests.post(API_URL, headers=headers, json=payload, timeout=60)
        
        # Log all API responses to console for debugging
        print(f"[API CALL] Status: {response.status_code}, Chat ID: {chat_id}, Model: {model}")
        print(f"[API CALL] Response headers: {dict(response.headers)}")
        
        if response.status_code == 200:
            result = response.json()
            
            # Log full response to console
            print(f"[API RESPONSE] Full response: {json.dumps(result, ensure_ascii=False, indent=2)[:1000]}...")
            
            # Extract content more safely
            choices = result.get('choices', [])
            if not choices or len(choices) == 0:
                ai_content = "Извините, не удалось получить ответ. (Нет choices в ответе)"
                print(f"[API ERROR] No choices in response: {result}")
            else:
                message = choices[0].get('message', {})
                ai_content = message.get('content', '') or ''
                
                # Remove reasoning tags from AI response
                if ai_content:
                    ai_content = clean_reasoning_tags(ai_content)
                
                # Log extracted content
                print(f"[API RESPONSE] Extracted content length: {len(ai_content)}, Preview: {ai_content[:200]}...")
            
            # Handle empty content - don't return error, just log warning
            if not ai_content or not ai_content.strip():
                print(f"[API WARNING] Empty content after processing - this may be normal for some models")
                ai_content = ""  # Keep empty, let frontend handle or check DB
            
            # Save to DB IMMEDIATELY after receiving response (even if client fails to receive)
            with sqlite3.connect(DB_FILE) as conn:
                cursor = conn.cursor()
                cursor.execute('INSERT INTO messages (chat_id, role, content) VALUES (?, ?, ?)', 
                               (chat_id, 'assistant', ai_content))
                conn.commit()
                print(f"[DB] Saved assistant message to DB for chat {chat_id}")
            
            # Analyze and store memory - return new entries in response
            new_memory_entries = []
            if memory_enabled and user_message:
                import threading
                memory_result = {'entries': []}
                def analyze_memory_background():
                    try:
                        new_entries = analyze_and_store_memory(user_id, user_message, ai_content)
                        memory_result['entries'] = new_entries
                    except Exception as e:
                        print(f"[MEMORY] Background analysis error: {e}")
                
                # Start analysis in background
                thread = threading.Thread(target=analyze_memory_background, daemon=True)
                thread.start()
                thread.join(timeout=3)  # Wait up to 3 seconds for memory analysis
                new_memory_entries = memory_result.get('entries', []) or []  # Ensure it's always a list, never None
            
            # For new chats with images, generate title AFTER AI response
            # so we can use the AI's description of the image
            if is_new_chat and images and len(images) > 0:
                import threading
                def generate_title_after_response():
                    try:
                        print(f"[TITLE GEN] Starting title generation AFTER response for chat {chat_id} with images")
                        # Use AI's response to understand what's in the image
                        # Extract first few sentences from AI response for context
                        ai_context = ai_content[:300] if len(ai_content) > 300 else ai_content
                        
                        # Generate title based on AI's response about the image
                        title_prompt = f"""На основе следующего описания изображения, придумай краткое название (2-4 слова) для чата. 
                        
Описание: {ai_context}

Только название на русском, без кавычек."""
                        
                        # Use vision model to generate title
                        fast_models = [
                            "Qwen/Qwen2.5-VL-32B-Instruct",  # Vision model
                            "mistralai/Mistral-Large-Instruct-2411",  # Fast fallback
                            "deepseek-ai/DeepSeek-V3.2",
                        ]
                        
                        ai_title = None
                        for model_name in fast_models:
                            try:
                                response_title = requests.post(API_URL, headers={
                                    "Content-Type": "application/json",
                                    "Authorization": f"Bearer {API_TOKEN}"
                                }, json={
                                    "model": model_name,
                                    "messages": [
                                        {"role": "user", "content": title_prompt}
                                    ],
                                    "temperature": 0.5,
                                    "max_tokens": 20
                                }, timeout=10)
                                
                                if response_title.status_code == 200:
                                    result_title = response_title.json()
                                    if 'choices' in result_title and len(result_title['choices']) > 0:
                                        content_title = result_title['choices'][0]['message'].get('content')
                                        if content_title:
                                            title = clean_reasoning_tags(content_title)
                                            title = title.strip().strip('"').strip("'").strip('«').strip('»')
                                            title = title.split('\n')[0].strip()
                                            if title and len(title) <= 50 and not '<' in title:
                                                ai_title = title
                                                break
                            except Exception as e:
                                print(f"[TITLE GEN] Model {model_name} failed: {e}")
                                continue
                        
                        if ai_title:
                            with sqlite3.connect(DB_FILE) as conn:
                                cursor = conn.cursor()
                                cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (ai_title, chat_id))
                                conn.commit()
                                print(f"[TITLE GEN] Title updated to: '{ai_title}' for chat {chat_id} (after AI response)")
                    except Exception as e:
                        print(f"[TITLE GEN] Error generating title after response: {e}")
                        import traceback
                        traceback.print_exc()
                
                # Start title generation after response in background
                threading.Thread(target=generate_title_after_response, daemon=True).start()
                
            # Always return valid response, even if content is empty
            # Frontend will check DB if content is empty
            response_data = {
                'role': 'assistant', 
                'content': ai_content if ai_content else '',
                'memory_updated': len(new_memory_entries) > 0,
                'new_memory': new_memory_entries  # List of new memory entries
            }
            
            print(f"[API SUCCESS] Returning response to client for chat {chat_id} (content length: {len(ai_content) if ai_content else 0})")
            return jsonify(response_data)
        else:
            error_msg = f"External API Error: {response.status_code}"
            error_details = response.text[:500]  # Limit error text length
            print(f"[API ERROR] {error_msg}: {error_details}")
            return jsonify({'error': error_msg, 'details': error_details}), 502

    except requests.exceptions.Timeout as e:
        error_msg = "Timeout waiting for API response"
        print(f"[API ERROR] {error_msg}: {e}")
        return jsonify({'error': error_msg}), 504
    except requests.exceptions.RequestException as e:
        error_msg = f"Network error: {str(e)}"
        print(f"[API ERROR] {error_msg}")
        return jsonify({'error': error_msg}), 502
    except Exception as e:
        error_msg = f"Unexpected error: {str(e)}"
        print(f"[API ERROR] {error_msg}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': error_msg}), 500

# ... (Profile routes and others same as original) ...
@app.route('/api/profile', methods=['GET'])
@login_required
def get_profile():
    user_id = str(session['user'].get('id'))
    try:
        with sqlite3.connect(DB_FILE) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute('SELECT * FROM user_profiles WHERE user_id = ?', (user_id,))
            profile = cursor.fetchone()
            
            # Get user memory
            cursor.execute('SELECT id, memory_text, created_at FROM user_memory WHERE user_id = ? ORDER BY created_at DESC', (user_id,))
            memory_rows = cursor.fetchall()
            memories = [{'id': r['id'], 'text': r['memory_text'], 'created_at': r['created_at']} for r in memory_rows]
            
            if profile:
                profile_dict = dict(profile)
                profile_dict['memory'] = memories
                # Ensure boolean values
                profile_dict['memory_enabled'] = bool(profile_dict.get('memory_enabled', True))
                profile_dict['notifications_enabled'] = bool(profile_dict.get('notifications_enabled', False))
                
                # Validate and clean avatar_url
                avatar_url = profile_dict.get('avatar_url', '')
                if avatar_url and (avatar_url == 'url' or not avatar_url.startswith(('http', '/static/'))):
                    # Invalid avatar_url - clear it
                    profile_dict['avatar_url'] = ''
                
                return jsonify(profile_dict)
            else:
                # Return defaults from session
                # Get preferred model from profile or use dreamGPT as default
                preferred = None
                cursor.execute('SELECT preferred_model FROM user_profiles WHERE user_id = ?', (user_id,))
                pref_row = cursor.fetchone()
                if pref_row and pref_row[0]:
                    preferred = pref_row[0]
                
                return jsonify({
                    'custom_name': session['user'].get('username', 'User'),
                    'avatar_url': '',
                    'system_prompt': 'Ты полезный ассистент.',
                    'preferred_model': preferred or 'dreamgpt-ai',
                    'memory_enabled': True,
                    'notifications_enabled': False,
                    'memory': []
                })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/profile', methods=['POST'])
@login_required
def update_profile():
    user_id = str(session['user'].get('id'))
    data = request.json
    
    custom_name = data.get('custom_name')
    avatar_url = data.get('avatar_url')
    system_prompt = data.get('system_prompt')
    preferred_model = data.get('preferred_model')
    prompt_template = data.get('prompt_template')
    memory_enabled = data.get('memory_enabled')
    notifications_enabled = data.get('notifications_enabled')
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            
            # Check if profile exists
            cursor.execute('SELECT user_id FROM user_profiles WHERE user_id = ?', (user_id,))
            exists = cursor.fetchone()
            
            if exists:
                # Update existing profile
                updates = []
                values = []
                if custom_name is not None:
                    updates.append('custom_name=?')
                    values.append(custom_name)
                if avatar_url is not None:
                    # Validate avatar_url - must be valid URL or empty string
                    if avatar_url and avatar_url.strip() and avatar_url != 'url':
                        # Only update if it's a valid URL or local path
                        if avatar_url.startswith('http') or avatar_url.startswith('/static/'):
                            updates.append('avatar_url=?')
                            values.append(avatar_url)
                        # Otherwise skip (invalid URL)
                    elif not avatar_url or avatar_url.strip() == '':
                        # Allow clearing avatar
                        updates.append('avatar_url=?')
                        values.append('')
                if system_prompt is not None:
                    updates.append('system_prompt=?')
                    values.append(system_prompt)
                if preferred_model is not None:
                    updates.append('preferred_model=?')
                    values.append(preferred_model)
                if prompt_template is not None:
                    updates.append('prompt_template=?')
                    values.append(prompt_template)
                if memory_enabled is not None:
                    updates.append('memory_enabled=?')
                    values.append(1 if memory_enabled else 0)
                if notifications_enabled is not None:
                    updates.append('notifications_enabled=?')
                    values.append(1 if notifications_enabled else 0)
                
                if updates:
                    updates.append('updated_at=CURRENT_TIMESTAMP')
                    values.append(user_id)
                    cursor.execute(f'''
                        UPDATE user_profiles SET {", ".join(updates)}
                        WHERE user_id=?
                    ''', values)
            else:
                # Insert new profile
                cursor.execute('''
                    INSERT INTO user_profiles (user_id, custom_name, avatar_url, system_prompt, preferred_model, 
                                             prompt_template, memory_enabled, notifications_enabled)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    user_id,
                    custom_name if custom_name is not None else None,
                    avatar_url if avatar_url is not None else None,
                    system_prompt if system_prompt is not None else None,
                    preferred_model if preferred_model is not None else None,
                    prompt_template if prompt_template is not None else None,
                    1 if memory_enabled is None else (1 if memory_enabled else 0),
                    0 if notifications_enabled is None else (1 if notifications_enabled else 0)
                ))
            conn.commit()
        return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/profile/avatar', methods=['POST'])
@login_required
def upload_avatar():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        user_id = str(session['user'].get('id'))
        filename = f"{user_id}_{uuid.uuid4().hex[:8]}.{file.filename.rsplit('.', 1)[1].lower()}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        avatar_url = f"/static/uploads/avatars/{filename}"
        
        # Update profile
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('''
                INSERT INTO user_profiles (user_id, avatar_url)
                VALUES (?, ?)
                ON CONFLICT(user_id) DO UPDATE SET
                avatar_url=excluded.avatar_url,
                updated_at=CURRENT_TIMESTAMP
            ''', (user_id, avatar_url))
            conn.commit()
        
        return jsonify({'avatar_url': avatar_url})
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/static/uploads/avatars/<filename>')
def uploaded_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename)

@app.route('/api/upload/image', methods=['POST'])
@login_required
def upload_chat_image():
    """Upload image for chat messages"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and allowed_file(file.filename):
        user_id = str(session['user'].get('id'))
        filename = f"{user_id}_{uuid.uuid4().hex[:8]}_{secure_filename(file.filename)}"
        filepath = os.path.join(CHAT_IMAGES_FOLDER, filename)
        file.save(filepath)
        
        image_url = f"/static/uploads/chat_images/{filename}"
        return jsonify({'url': image_url, 'filename': filename})
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/api/upload/document', methods=['POST'])
@login_required
def upload_chat_document():
    """Upload document file (PDF, TXT, DOCX, etc.) or VIDEO for chat messages"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Check allowed extensions (docs + video)
    is_doc = allowed_document_file(file.filename)
    is_video = allowed_video_file(file.filename)
    
    if file and (is_doc or is_video):
        user_id = str(session['user'].get('id'))
        filename = f"{user_id}_{uuid.uuid4().hex[:8]}_{secure_filename(file.filename)}"
        filepath = os.path.join(CHAT_DOCUMENTS_FOLDER, filename)
        file.save(filepath)
        
        document_url = f"/static/uploads/chat_documents/{filename}"
        return jsonify({'url': document_url, 'filename': filename, 'type': file.content_type})
    
    return jsonify({'error': 'Invalid file type. Supported: PDF, TXT, DOCX, DOC, RTF, ODT, MP4, MOV, AVI'}), 400

@app.route('/static/uploads/chat_images/<filename>')
def chat_image_file(filename):
    return send_from_directory(CHAT_IMAGES_FOLDER, filename)

@app.route('/static/uploads/chat_documents/<filename>')
def chat_document_file(filename):
    return send_from_directory(CHAT_DOCUMENTS_FOLDER, filename)

@app.route('/api/chats/<chat_id>/title', methods=['POST'])
@login_required
def generate_chat_title_endpoint(chat_id):
    """Generate chat title after first message"""
    user_id = str(session['user'].get('id'))
    
    try:
        # Verify ownership
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('SELECT user_id FROM chats WHERE id = ?', (chat_id,))
            row = cursor.fetchone()
            
            if not row or row[0] != user_id:
                return jsonify({'error': 'Access denied'}), 403
            
            # Get first user message
            cursor.execute('SELECT content FROM messages WHERE chat_id = ? AND role = ? ORDER BY id ASC LIMIT 1', 
                          (chat_id, 'user'))
            msg_row = cursor.fetchone()
            
            if not msg_row:
                return jsonify({'error': 'No messages found'}), 404
            
            user_message = msg_row[0]
            
            # Check if user_message is valid
            if not user_message or not isinstance(user_message, str):
                user_message = "Новый чат"
            
            # Generate title
            ai_title = generate_chat_title(user_message)
            if not ai_title or not isinstance(ai_title, str):
                # Fallback: use first word(s) of message
                if user_message and user_message.strip():
                    words = user_message.strip().split()
                    ai_title = words[0] if words else "Новый чат"
                    if len(words) > 1:
                        ai_title = " ".join(words[:2])
                else:
                    ai_title = "Новый чат"
            
            # Update title
            cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (ai_title, chat_id))
            conn.commit()
            
            return jsonify({'title': ai_title})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chat/topics')
@login_required
def get_topics():
    """Get 4 cached welcome topics from database"""
    try:
        with sqlite3.connect(DB_FILE) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            # Get random 4 topics from cached set
            cursor.execute('SELECT title, description, prompt FROM welcome_topics ORDER BY RANDOM() LIMIT 4')
            rows = cursor.fetchall()
            
            if len(rows) < 4:
                # Not enough topics in DB, use fallback
                selected = random.sample(ALL_TOPICS, 4)
                return jsonify([{'title': t['title'], 'desc': t['desc'], 'prompt': t['prompt']} for t in selected])
            
            topics = [{'title': r['title'], 'desc': r['description'], 'prompt': r['prompt']} for r in rows]
            return jsonify(topics)
    except Exception as e:
        print(f"Error loading topics: {e}")
        # Fallback to predefined
        selected = random.sample(ALL_TOPICS, 4)
        return jsonify([{'title': t['title'], 'desc': t['desc'], 'prompt': t['prompt']} for t in selected])

# Memory endpoints
@app.route('/api/memory', methods=['GET'])
@login_required
def get_memory():
    """Get all memory entries for the user"""
    user_id = str(session['user'].get('id'))
    try:
        with sqlite3.connect(DB_FILE) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            cursor.execute('SELECT id, memory_text, created_at FROM user_memory WHERE user_id = ? ORDER BY created_at DESC', (user_id,))
            rows = cursor.fetchall()
            memories = [{'id': r['id'], 'text': r['memory_text'], 'created_at': r['created_at']} for r in rows]
            return jsonify(memories)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory', methods=['POST'])
@login_required
def add_memory():
    """Add a new memory entry"""
    user_id = str(session['user'].get('id'))
    data = request.json
    memory_text = data.get('text', '').strip()
    
    if not memory_text:
        return jsonify({'error': 'Memory text is required'}), 400
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            cursor.execute('INSERT INTO user_memory (user_id, memory_text) VALUES (?, ?)', (user_id, memory_text))
            conn.commit()
            return jsonify({'id': cursor.lastrowid, 'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/<int:memory_id>', methods=['PUT'])
@login_required
def update_memory(memory_id):
    """Update a memory entry"""
    user_id = str(session['user'].get('id'))
    data = request.json
    memory_text = data.get('text', '').strip()
    
    if not memory_text:
        return jsonify({'error': 'Memory text is required'}), 400
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            # Verify ownership
            cursor.execute('SELECT user_id FROM user_memory WHERE id = ?', (memory_id,))
            row = cursor.fetchone()
            if not row or row[0] != user_id:
                return jsonify({'error': 'Access denied'}), 403
            
            cursor.execute('UPDATE user_memory SET memory_text = ? WHERE id = ?', (memory_text, memory_id))
            conn.commit()
            return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/memory/<int:memory_id>', methods=['DELETE'])
@login_required
def delete_memory(memory_id):
    """Delete a memory entry"""
    user_id = str(session['user'].get('id'))
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            # Verify ownership
            cursor.execute('SELECT user_id FROM user_memory WHERE id = ?', (memory_id,))
            row = cursor.fetchone()
            if not row or row[0] != user_id:
                return jsonify({'error': 'Access denied'}), 403
            
            cursor.execute('DELETE FROM user_memory WHERE id = ?', (memory_id,))
            conn.commit()
            return jsonify({'status': 'success'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/chats/<chat_id>/title', methods=['PUT'])
@login_required
def update_chat_title(chat_id):
    """Update chat title manually"""
    user_id = str(session['user'].get('id'))
    data = request.json
    new_title = data.get('title', '').strip()
    
    if not new_title:
        return jsonify({'error': 'Title is required'}), 400
    
    try:
        with sqlite3.connect(DB_FILE) as conn:
            cursor = conn.cursor()
            # Verify ownership
            cursor.execute('SELECT user_id FROM chats WHERE id = ?', (chat_id,))
            row = cursor.fetchone()
            if not row or row[0] != user_id:
                return jsonify({'error': 'Access denied'}), 403
            
            cursor.execute('UPDATE chats SET title = ? WHERE id = ?', (new_title, chat_id))
            conn.commit()
            return jsonify({'status': 'success', 'title': new_title})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Production port for nginx proxy
    app.run(debug=False, host='127.0.0.1', port=5027)
